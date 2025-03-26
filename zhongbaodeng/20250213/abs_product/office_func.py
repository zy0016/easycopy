import sys
import docx
from docx import Document
from docx.shared import RGBColor
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.oxml import OxmlElement
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from openpyxl.styles import Alignment
from openpyxl.styles import PatternFill
from openpyxl.styles import Border, Side, Font
try:
    from abs_product.func import savelog, get_row, get_max_row_maps, get_list_max_value, get_format_value, update_tup, merge_maps_data
except ImportError as e:
    from func import savelog, get_row, get_max_row_maps, get_list_max_value, get_format_value, update_tup, merge_maps_data


excel_alphabet=['','A','B','C','D','E','F','G','H','I','J','K','L','M','N','O','P','Q','R','S','T','U','V','W','X','Y','Z']
excel_border_style = Border(
            top=Side(border_style='thin', color='000000'),
            bottom=Side(border_style='thin', color='000000'),
            left=Side(border_style='thin', color='000000'),
            right=Side(border_style='thin', color='000000')
        )
# 填充表格内容并设置背景颜色
# 第一行设置背景颜色
def set_cell_background(cell, color):
    cell._element.get_or_add_tcPr().append(
        docx.oxml.parse_xml(r'<w:shd {} w:fill="{}"/>'.format(
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"',
            color)
        )
    )

def set_cell_font_color(cell, color):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = color

def set_font(paragraph, font_name, font_size):
    run2 = paragraph.runs[0]
    run2.font.name = font_name
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run2.font.size = Pt(font_size)

def set_text_align(cell,horizontal=True,vertical=True):
    if horizontal:
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if vertical:
        cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

def set_row_height(table, row_idx, height):
    tr = table.rows[row_idx]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

def create_word_table():
    #创建word文件，在其中创建表格
    # 创建一个新的Document对象
    doc = Document()
    # 添加标题
    doc.add_heading('Python创建表格示例', 0)
    doc.add_paragraph('这是一个简单的示例，展示如何使用 python 创建并编辑 Word 文件，以及在word中创建表格。')
    para = doc.add_paragraph()
    runt1 = para.add_run('前面的文字使用默认颜色')
    runt2 = para.add_run('后面的文字使用红色字体')
    runt2.font.color.rgb = RGBColor(255,0,0)
    # 添加一个表格，表格大小为3行4列
    table = doc.add_table(rows=3, cols=5)
    
    set_cell_background(table.cell(0, 0), 'D9D9D9') 
    set_cell_background(table.cell(0, 1), 'D9D9D9')
    set_cell_background(table.cell(0, 2), 'D9D9D9')
    set_cell_background(table.cell(0, 3), 'D9D9D9')
    set_cell_background(table.cell(0, 4), 'D9D9D9')
    table.style = 'Table Grid'
    # 获取表格的单元格并填写内容
    # 填写标题
    cell_00 = table.cell(0, 0)
    cell_00.text = '序\n号'
    set_text_align(cell_00)
    set_font(cell_00.paragraphs[0],'仿宋_GB2312',11)
    
    cell_01 = table.cell(0, 1)
    cell_01.text = '产品名称'
    set_text_align(cell_01)
    set_font(cell_01.paragraphs[0],'仿宋_GB2312',11)

    cell_02 = table.cell(0, 2)
    cell_02.text = '工作进度'
    set_text_align(cell_02)
    set_font(cell_02.paragraphs[0],'仿宋_GB2312',11)
    
    cell_03 = table.cell(0, 3)
    cell_03.text = '产品规模(亿元)'
    set_text_align(cell_03)
    set_font(cell_03.paragraphs[0],'仿宋_GB2312',11)

    cell_04 = table.cell(0, 4)
    cell_04.text = '备注'
    set_text_align(cell_04)
    set_font(cell_04.paragraphs[0],'仿宋_GB2312',11)
    ################填写内容#####################################################
    cell_10 = table.cell(1, 0)
    cell_10.text = 'x'
    set_text_align(cell_10)
    
    cell_11 = table.cell(1, 1)
    paragraph = cell_11.paragraphs[0]
    run1 = paragraph.add_run('xx')
    run1.font.color.rgb = RGBColor(0,0,0)
    run2 = paragraph.add_run('（资产支持计划名称）')
    run2.font.color.rgb = RGBColor(255,0,0)
    set_text_align(cell_11,False,True)

    cell_12 = table.cell(1, 2)
    #cell_12.text = '1.xx年x月x日报送系统，当日分配人员进行查验。\n2.xx年x月x日xx（予以登记/不予登记)。'
    paragraph = cell_12.paragraphs[0]
    run121 = paragraph.add_run('1.xx年x月x日报送系统，当日分配人员进行查验。\n2.xx年x月x日xx')
    run121.font.color.rgb = RGBColor(0,0,0)
    run122 = paragraph.add_run('（予以登记/不予登记）')
    run122.font.color.rgb = RGBColor(255,0,0)
    run123 = paragraph.add_run('。')
    run123.font.color.rgb = RGBColor(0,0,0)
    
    cell_13 = table.cell(1, 3)
    cell_13.text = 'x'
    set_text_align(cell_13)
    
    cell_14 = table.cell(1, 4)
    cell_14.text = '予以登记/不予登记/查验中。'
    set_cell_font_color(cell_14, RGBColor(255, 0, 0)) 
    set_text_align(cell_14,False,True)

    #设置单元格宽度
    for row in table.rows:
        cell0 = row.cells[0]
        cell0.width = Inches(0.5)
        cell2 = row.cells[2]
        cell2.width = Inches(2.5)

    #设置标题栏高度
    set_row_height(table,0,1000)

    # 保存文档
    doc.save('word_table.docx')

    print("Word文件已生成！")

def get_row_line(ws,bottom_line):
    #获得合适的可以显示文字的行号
    while bottom_line > 0:
        if ws.cell(row=bottom_line,column=1).value is not None:
            break
        bottom_line = bottom_line - 1

    return bottom_line + 3

def set_excel_sheet_map_text(ws,row_start,col_start,map1,map2,write_name = True):
    alphabet=['','A','B','C','D','E','F','G','H','I','J','K','L','M','N','O','P','Q','R','S','T','U','V','W','X','Y','Z']
    line = row_start
    id1 = 0
    id2 = 0
    id3 = 0
    id4 = 0
    if write_name:
        id1 = col_start
        id2 = col_start + 2
        id3 = col_start + 2
        id4 = col_start + 3
    else:
        id1 = col_start - 1
        id2 = col_start + 1
        id3 = col_start + 1
        id4 = col_start + 2
    
    for k1 in map1:
        r1 = map1[k1][1]
        v1 = map1[k1][0]
        ws.cell(row=line + r1,column=1).value = k1

        datatoday = alphabet[id4 - 2]
        datapre = alphabet[id4 - 1]
        item = "=IF(" + datapre + str(line + r1) + "=0,\"\",(" + datatoday + str(line + r1) + "-" + datapre + str(line + r1) + ")/" + datapre + str(line + r1) + ")"
        ws.cell(row=line + r1,column=id4).value = item
        ws.cell(row=line + r1,column=id4).number_format = '0.00%'

        if ws.cell(row=line + r1,column=id4-2).value is not None and ws.cell(row=line + r1,column=id4-1).value is not None:
            ws.cell(row=line + r1,column=id4).value = item

        ws.cell(row=line + r1,column=id1 + 1).value = get_format_value(v1)

    map3 = {}
    for k2 in map2:
        if k2 in map1:
            r2 = get_row(map1,k2)
            v2 = map2[k2][0]
            ws.cell(row=line + r2,column=id2).value = get_format_value(v2)
        else:
            map3[k2] = map2[k2]

    id1 = col_start
    for k3 in map3:
        r3 = map3[k3][1]
        v3 = map3[k3][0]
        if ws.cell(row=line + r3,column=1).value is None or len(ws.cell(row=line + r3,column=1).value) == 0:
            ws.cell(row=line + r3,column=1).value = k3

        datatoday = alphabet[id4 - 2]
        datapre = alphabet[id4 - 1]
        item = "=IF(" + datapre + str(line + r3) + "=0,\"\",(" + datatoday + str(line + r3) + "-" + datapre + str(line + r3) + ")/" + datapre + str(line + r3) + ")"
        ws.cell(row=line + r3,column=id4).value = item
        ws.cell(row=line + r3,column=id4).number_format = '0.00%'
        
        if ws.cell(row=line + r3,column=id4-2).value is not None and ws.cell(row=line + r3,column=id4-1).value is not None:
            ws.cell(row=line + r3,column=id4).value = item

        ws.cell(row=line + r3,column=id3).value = get_format_value(v3)

    return get_max_row_maps(map1,map2) + line + 1

def set_excel_sheet_row_text(ws,line,text_list):
    #设置excel文件某一行的文字
    #如果column为零,则在这一行显示所有数据。否则text_list将在指定列换行显示，函数返回当前显示的最后一行的行号
    border_style = Border(
        top=Side(border_style='thin', color='000000'),
        bottom=Side(border_style='thin', color='000000'),
        left=Side(border_style='thin', color='000000'),
        right=Side(border_style='thin', color='000000')
    )
    id = 1
    alphabet=['A','B','C','D','E','F','G','H','I','J','K','L','M','N','O','P','Q','R','S','T','U','V','W','X','Y','Z']
    if len(text_list) > len(alphabet):
        print('set_excel_sheet_row_text text_list and alphabet are not match!')
        sys.exit(1)

    for item in text_list:
        if item == "calc":
            datatoday = alphabet[id - 3]
            datapre = alphabet[id - 2]
            item = "=IF(" + datapre + str(line) + "=0,\"\",(" + datatoday + str(line) + "-" + datapre + str(line) + ")/" + datapre + str(line) + ")"
            ws.cell(row=line,column=id).number_format = '0.00%'

        ws.cell(row=line,column=id).value = item
        cell_id = alphabet[id - 1]
        ws.column_dimensions[cell_id].width = 15
        ws[cell_id + str(line)].border = border_style
        if len(str(item)) <= 6:
            ws[cell_id + str(line)].alignment = Alignment(vertical='center')
        else:
            ws[cell_id + str(line)].alignment = Alignment(wrap_text=True)

        id = id + 1

    return line + 1

def set_excel_sheet_row_text_from_col(ws,line,col,text_list,fill_color = None,font = None):
    for item in text_list:
        ws.cell(row=line,column=col).value = item
        cell_id = excel_alphabet[col]
        ws.column_dimensions[cell_id].width = 15
        if fill_color is not None:
            ws[cell_id + str(line)].fill = fill_color
        if font is not None:
            ws[cell_id + str(line)].font = font
        ws[cell_id + str(line)].border = excel_border_style
        if len(str(item)) <= 6:
            ws[cell_id + str(line)].alignment = Alignment(horizontal='center', vertical='center')
        else:
            ws[cell_id + str(line)].alignment = Alignment(wrap_text=True)

        col = col + 1

def set_merge_cells_border_row(ws,bstart,bend):
        #设置合并后的单元格边框
        start = bstart[0]
        end = bend[0]
        row = str(int(bstart[1:]))
        i = 1
        bstart = False
        while i < len(excel_alphabet):
            if start == excel_alphabet[i]:
                bstart = True
            if bstart and end == excel_alphabet[i]:
                bstart = False
                ws[excel_alphabet[i] + row].border = excel_border_style
            
            if bstart == True:
                ws[excel_alphabet[i] + row].border = excel_border_style

            i = i + 1

def set_month(ws,line,id,font1):
    for i in range(12):
        ws.cell(row=line,column=id).value = str(i + 1) + "月"
        cell_id = excel_alphabet[id]
        ws[cell_id + str(line)].border = excel_border_style
        ws[cell_id + str(line)].font = font1
        ws[cell_id + str(line)].alignment = Alignment(horizontal='center', vertical='center')
        line = line + 1

    ws.cell(row=line,column=id).value = '总计'
    cell_id = excel_alphabet[id]
    ws[cell_id + str(line)].border = excel_border_style
    ws[cell_id + str(line)].font = font1
    ws[cell_id + str(line)].alignment = Alignment(horizontal='center', vertical='center')

def set_excel_asset_support_plan(ws,line,col,datamap):
    #设置excel文件的"资产支持计划登记情况"内容
    for k in datamap:
        lst = [str(k),\
                datamap[k]["C_FUNDNAME"],\
                datamap[k]["mem_name"],\
                datamap[k]["F_PRDREGBALA"],\
                datamap[k]["C_INFRACLASS"],\
                datamap[k]["d_preregisterdate"],\
                datamap[k]["days_between"],\
                datamap[k]["reportdate"],\
                datamap[k]["C_AUDITFLAG"]]
        set_excel_sheet_row_text_from_col(ws,line,col,lst)
        line = line + 1

def set_excel_reg_group_asset(ws,line,col,datamap):
    #设置excel文件中 “予以登记”组合资管产品明细
    for k in datamap:
        lst = [str(k),datamap[k]["C_FUNDNAME"],datamap[k]["C_PRDTYPE"],datamap[k]["C_OPERATEWAY"],datamap[k]["C_REGISTERCLASSIY"],datamap[k]["C_REG_PROGRESS"]]
        set_excel_sheet_row_text_from_col(ws,line,col,lst)
        line = line + 1
    return line

def set_product_reg_type(ws,line,col,datamap):
    #设置excel文件中一年十二个月的产品类型（只）
    v0 = 0
    v1 = 0
    v2 = 0
    v9 = 0
    for i in range(13):
        if i == 0:
            continue
        vmap = datamap[str(i)]
        # 0:固定收益类
        # 1:权益类
        # 2:商品及金融衍生品类
        # 9:混合类
        lst = [int(vmap["0"]),int(vmap["1"]),int(vmap["2"]),int(vmap["9"])]
        v0 = v0 + int(vmap["0"])
        v1 = v1 + int(vmap["1"])
        v2 = v2 + int(vmap["2"])
        v9 = v9 + int(vmap["9"])
        set_excel_sheet_row_text_from_col(ws,line,col,lst)
        line = line + 1

    set_excel_sheet_row_text_from_col(ws,line,col,[v0,v1,v2,v9])
    return line

def set_working_type(ws,line,col,datamap):
    #设置excel文件中一年十二个月的运作方式（只）
    v1 = 0
    v2 = 0
    for i in range(13):
        if i == 0:
            continue
        vmap = datamap[str(i)]
        # 1:开放式
        # 2:封闭式
        lst = [int(vmap["1"]),int(vmap["2"])]
        v1 = v1 + int(vmap["1"])
        v2 = v2 + int(vmap["2"])
        set_excel_sheet_row_text_from_col(ws,line,col,lst)
        line = line + 1

    set_excel_sheet_row_text_from_col(ws,line,col,[v1,v2])
    return line

def set_manager_type(ws,line,col,datamap):
    id = 1
    for key in datamap:
        vmap = datamap[key]
        lst = [str(id),key,vmap["fixed"],vmap["right"],vmap["product_finance"],vmap["mixed"],vmap["open"],vmap["close"]]
        set_excel_sheet_row_text_from_col(ws,line,col,lst)
        id = id + 1
        line = line + 1
    return line

def set_merge_cells_border_col(ws,bstart,bend):
    cell = bstart[0]
    start = int(bstart[1:])
    end = int(bend[1:])
    while start <= end:
        ws[cell + str(start)].border = excel_border_style
        start = start + 1


def create_excel_file(xls_data):
    #创建excel文件内容
    wb = Workbook()
    fill_green = PatternFill(start_color="92D050", end_color="92D050", fill_type="solid") 
    fill_yellow = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
    # 设置颜色
    # font_style = Font(color='FF0000')
    #ws1['M2'].font = font_style
    ws1 = wb.active
    #############################################################
    ws1.title = '月度'
    ws1.merge_cells('A1:M1')
    ws1['A1'] = '中保登资产支持计划登记情况表（月度）'
    ws1['A1'].alignment = Alignment(horizontal='center', vertical='center')

    line = 2
    set_excel_sheet_row_text(ws1,line,xls_data.title_month_total_lst)
    ws1['A' + str(line)].fill = fill_green
    ws1['H' + str(line)].fill = fill_yellow

    line = 3
    set_excel_sheet_row_text(ws1,line,xls_data.month_total_lst)
    ws1['H' + str(line)].fill = fill_yellow

    line = 6
    line = set_excel_sheet_row_text(ws1,line,xls_data.title_month_trustee_lst)
    ws1['A' + str(6)].fill = fill_green

    update_tup(xls_data.m_trustee)

    i1 = set_excel_sheet_map_text(ws1,7,1,xls_data.m_trustee[0],xls_data.m_trustee[1],True)
    i2 = set_excel_sheet_map_text(ws1,7,5,xls_data.m_trustee[2],xls_data.m_trustee[3],False)
    i3 = set_excel_sheet_map_text(ws1,7,8,xls_data.m_trustee[4],xls_data.m_trustee[5],False)
    i4 = set_excel_sheet_map_text(ws1,7,11,xls_data.m_trustee[6],xls_data.m_trustee[7],False)
    l1 = [i1,i2,i3,i4]
    
    line = get_list_max_value(l1)
    line = get_row_line(ws1,line)

    line = set_excel_sheet_row_text(ws1,line,xls_data.title_month_deposit_bank_lst)
    ws1['A' + str(line - 1)].fill = fill_green

    update_tup(xls_data.m_deposit_bank)
    i1 = set_excel_sheet_map_text(ws1,line,1,xls_data.m_deposit_bank[0],xls_data.m_deposit_bank[1],True)
    i2 = set_excel_sheet_map_text(ws1,line,5,xls_data.m_deposit_bank[2],xls_data.m_deposit_bank[3],False)
    i3 = set_excel_sheet_map_text(ws1,line,8,xls_data.m_deposit_bank[4],xls_data.m_deposit_bank[5],False)
    i4 = set_excel_sheet_map_text(ws1,line,11,xls_data.m_deposit_bank[6],xls_data.m_deposit_bank[7],False)
    ########################################################################################################
    ws2 = wb.create_sheet('季度')
    ws2.merge_cells('A1:M1')
    ws2['A1'] = '中保登资产支持计划登记情况表（季度）'
    ws2['A1'].alignment = Alignment(horizontal='center', vertical='center')

    line = 2
    set_excel_sheet_row_text(ws2,line,xls_data.title_quarter_total_lst)
    ws2['A' + str(line)].fill = fill_green
    ws2['H' + str(line)].fill = fill_yellow

    line = 3
    set_excel_sheet_row_text(ws2,line,xls_data.quarter_total_lst)

    line = 6
    set_excel_sheet_row_text(ws2,line,xls_data.title_quarter_trustee_lst)
    ws2['A' + str(line)].fill = fill_green

    update_tup(xls_data.q_trustee)
    i1 = set_excel_sheet_map_text(ws2,7,1,xls_data.q_trustee[0],xls_data.q_trustee[1],True)
    i2 = set_excel_sheet_map_text(ws2,7,5,xls_data.q_trustee[2],xls_data.q_trustee[3],False)
    i3 = set_excel_sheet_map_text(ws2,7,8,xls_data.q_trustee[4],xls_data.q_trustee[5],False)
    i4 = set_excel_sheet_map_text(ws2,7,11,xls_data.q_trustee[6],xls_data.q_trustee[7],False)

    l1 = [i1,i2,i3,i4]
    
    line = get_list_max_value(l1)
    line = get_row_line(ws2,line)
    line = set_excel_sheet_row_text(ws2,line,xls_data.title_quarter_deposit_bank_lst)
    ws2['A' + str(line - 1)].fill = fill_green

    update_tup(xls_data.q_deposit_bank)
    
    i1 = set_excel_sheet_map_text(ws2,line,1,xls_data.q_deposit_bank[0],xls_data.q_deposit_bank[1],True)
    i2 = set_excel_sheet_map_text(ws2,line,5,xls_data.q_deposit_bank[2],xls_data.q_deposit_bank[3],False)
    i3 = set_excel_sheet_map_text(ws2,line,8,xls_data.q_deposit_bank[4],xls_data.q_deposit_bank[5],False)
    i4 = set_excel_sheet_map_text(ws2,line,11,xls_data.q_deposit_bank[6],xls_data.q_deposit_bank[7],False)
    ##########################################################################
    ws3 = wb.create_sheet('年度')
    ws3.merge_cells('A1:M1')
    ws3['A1'] = '中保登资产支持计划登记情况表（年度）'
    ws3['A1'].alignment = Alignment(horizontal='center', vertical='center')

    line = 2
    set_excel_sheet_row_text(ws3,line,xls_data.title_year_total_lst)
    ws3['A' + str(line)].fill = fill_green
    ws3['H' + str(line)].fill = fill_yellow

    line = 3
    set_excel_sheet_row_text(ws3,line,xls_data.year_total_lst)

    line = 6
    set_excel_sheet_row_text(ws3,line,xls_data.title_year_trustee_lst)
    ws3['A' + str(line)].fill = fill_green

    update_tup(xls_data.y_trustee)
    i1 = set_excel_sheet_map_text(ws3,7,1,xls_data.y_trustee[0],xls_data.y_trustee[1],True)
    i2 = set_excel_sheet_map_text(ws3,7,5,xls_data.y_trustee[2],xls_data.y_trustee[3],False)
    i3 = set_excel_sheet_map_text(ws3,7,8,xls_data.y_trustee[4],xls_data.y_trustee[5],False)
    i4 = set_excel_sheet_map_text(ws3,7,11,xls_data.y_trustee[6],xls_data.y_trustee[7],False)
    l1 = [i1,i2,i3,i4]

    line = get_list_max_value(l1)
    line = get_row_line(ws3,line)
    line = set_excel_sheet_row_text(ws3,line,xls_data.title_year_deposit_bank_lst)
    ws3['A' + str(line - 1)].fill = fill_green

    update_tup(xls_data.y_deposit_bank)
    i1 = set_excel_sheet_map_text(ws3,line,1,xls_data.y_deposit_bank[0],xls_data.y_deposit_bank[1],True)
    i2 = set_excel_sheet_map_text(ws3,line,5,xls_data.y_deposit_bank[2],xls_data.y_deposit_bank[3],False)
    i3 = set_excel_sheet_map_text(ws3,line,8,xls_data.y_deposit_bank[4],xls_data.y_deposit_bank[5],False)
    i4 = set_excel_sheet_map_text(ws3,line,11,xls_data.y_deposit_bank[6],xls_data.y_deposit_bank[7],False)
    ######################################################################################################################
    ws4 = wb.create_sheet('历史累计')
    ws4.merge_cells('A1:J1')
    ws4['A1'] = '中保登资产支持计划登记情况表（历史累计）'
    ws4['A1'].alignment = Alignment(horizontal='center', vertical='center')

    line = 2
    set_excel_sheet_row_text(ws4,line,xls_data.title_amount_total_lst)
    ws4['A' + str(line)].fill = fill_green
    ws4['H' + str(line)].fill = fill_yellow

    line = 3
    set_excel_sheet_row_text(ws4,line,xls_data.amount_total_lst)
    ws4['H' + str(line)].fill = fill_yellow

    line = 6
    set_excel_sheet_row_text(ws4,line,xls_data.title_amount_trustee_lst)
    ws4['A' + str(line)].fill = fill_green

    update_tup(xls_data.a_trustee)
    i1 = set_excel_sheet_map_text(ws4,7,1,xls_data.a_trustee[0],xls_data.a_trustee[1],True)
    i2 = set_excel_sheet_map_text(ws4,7,5,xls_data.a_trustee[2],xls_data.a_trustee[3],False)
    i3 = set_excel_sheet_map_text(ws4,7,8,xls_data.a_trustee[4],xls_data.a_trustee[5],False)
    l1 = [i1,i2,i3]

    line = get_list_max_value(l1)
    line = get_row_line(ws4,line)
    line = set_excel_sheet_row_text(ws4,line,xls_data.title_amount_deposit_bank_lst)
    ws4['A' + str(line - 1)].fill = fill_green

    update_tup(xls_data.a_deposit_bank)
    set_excel_sheet_map_text(ws4,line,1,xls_data.a_deposit_bank[0],xls_data.a_deposit_bank[1],True)
    set_excel_sheet_map_text(ws4,line,5,xls_data.a_deposit_bank[2],xls_data.a_deposit_bank[3],False)
    set_excel_sheet_map_text(ws4,line,8,xls_data.a_deposit_bank[4],xls_data.a_deposit_bank[5],False)
    ##########################################################################
    wb.active = ws1
    wb.save(xls_data.get_filepath())
    savelog(xls_data.get_filepath() + "文件已生成！")

def create_reg_month_excel_file(xls_data):
    #创建excel文件内容
    wb = Workbook()
    ws1 = wb.active
    ####################################################################
    ws1.title = xls_data.title1
    ws1.merge_cells('A1:I1')
    ws1['A1'] = xls_data.txt1
    ws1['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws1['A1'].fill = xls_data.fill_grey
    ws1['A1'].font = xls_data.font1
    set_excel_sheet_row_text_from_col(ws1,2,1,xls_data.lst1,xls_data.fill_grey,xls_data.font1)
    line = 3
    set_excel_asset_support_plan(ws1,line,1,xls_data.asset_support_plan_map)
    ###################################################################
    ws2 = wb.create_sheet(xls_data.title2)
    ws2.merge_cells('A1:F1')
    ws2['A1'] = xls_data.txt2
    ws2['A1'].font = xls_data.font1
    ws2['A1'].alignment = Alignment(horizontal='center', vertical='center')
    set_merge_cells_border_row(ws2,"A1","F1")
    line = 2
    set_excel_sheet_row_text_from_col(ws2,line,1,xls_data.lst2,None,xls_data.font1)
    line = 3
    line = set_excel_reg_group_asset(ws2,line,1,xls_data.reg_group_asset_map) + 2

    a1 = "A" + str(line)
    g1 = "G" + str(line)
    ws2.merge_cells(a1+":"+g1)
    ws2[a1] = xls_data.txt3
    ws2[a1].alignment = Alignment(horizontal='center', vertical='center')
    ws2[a1].font = xls_data.font1
    set_merge_cells_border_row(ws2,a1,g1)

    line = line + 1
    am1 = "A" + str(line)
    am2 = "A" + str(line + 1)
    ws2.merge_cells(am1 + ":" + am2)
    ws2[am1] = "月份"
    ws2[am1].font = xls_data.font1
    ws2[am1].alignment = Alignment(horizontal='center', vertical='center')
    ws2[am1].border = xls_data.border_style

    bm1 = "B" + str(line)
    em1 = "E" + str(line)
    ws2.merge_cells(bm1 + ":" + em1)
    ws2[bm1] = xls_data.txt4
    ws2[bm1].font = xls_data.font1
    ws2[bm1].alignment = Alignment(horizontal='center', vertical='center')
    set_merge_cells_border_row(ws2,bm1,em1)
    
    f1 = "F" + str(line)
    g2 = "G" + str(line)
    ws2.merge_cells(f1 + ":" + g2)
    ws2[f1] = xls_data.txt5
    ws2[f1].font = xls_data.font1
    ws2[f1].alignment = Alignment(horizontal='center', vertical='center')
    set_merge_cells_border_row(ws2,f1,g2)

    line = line + 1
    set_excel_sheet_row_text_from_col(ws2,line,2,xls_data.lst3,None,xls_data.font1)
    line = line + 1
    set_month(ws2,line,1,xls_data.font1)

    set_product_reg_type(ws2,line,2,xls_data.product_reg_type_map)
    line = set_working_type(ws2,line,6,xls_data.working_type_map) + 1

    line = line + 2
    am2 = "A" + str(line)
    im2 = "I" + str(line)
    ws2.merge_cells(am2 + ":" + im2)
    ws2[am2] = xls_data.txt3
    ws2[am2].font = xls_data.font1
    ws2[am2].alignment = Alignment(horizontal='center', vertical='center')
    set_merge_cells_border_row(ws2,am2,im2)

    line = line + 1
    amn1 = "A" + str(line)
    amn2 = "A" + str(line + 1)
    ws2.merge_cells(amn1 + ":" + amn2)
    ws2[amn1] = '序号'
    ws2[amn1].font = xls_data.font1
    ws2[amn1].alignment = Alignment(horizontal='center', vertical='center')
    set_merge_cells_border_col(ws2,amn1,amn2)

    bmn1 = "B" + str(line)
    bmn2 = "B" + str(line + 1)
    ws2.merge_cells(bmn1 + ":" + bmn2)
    ws2[bmn1] = '管理人'
    ws2[bmn1].font = xls_data.font1
    ws2[bmn1].alignment = Alignment(horizontal='center', vertical='center')
    set_merge_cells_border_col(ws2,bmn1,bmn2)

    c1 = 'C' + str(line)
    f1 = "F" + str(line)
    ws2.merge_cells(c1 + ":" + f1)
    ws2[c1] = '产品类型'
    ws2[c1].font = xls_data.font1
    ws2[c1].alignment = Alignment(horizontal='center', vertical='center')
    set_merge_cells_border_row(ws2,c1,f1)

    g1 = "G" + str(line)
    h1 = "H" + str(line)
    ws2.merge_cells(g1 + ":" + h1)
    ws2[g1] = '运作方式'
    ws2[g1].font = xls_data.font1
    ws2[g1].alignment = Alignment(horizontal='center', vertical='center')
    set_merge_cells_border_row(ws2,g1,h1)

    i1 = "I" + str(line)
    i2 = "I" + str(line + 1)
    ws2.merge_cells(i1 + ":" + i2)
    ws2[i1] = '合计'
    ws2[i1].font = xls_data.font1
    ws2[i1].alignment = Alignment(horizontal='center', vertical='center')
    set_merge_cells_border_row(ws2,g1,h1)
    set_merge_cells_border_col(ws2,i1,i2)

    line = line + 1
    set_excel_sheet_row_text_from_col(ws2,line,3,xls_data.lst3,None,xls_data.font1)

    line = line + 1
    resmap = merge_maps_data(xls_data.manager_product_type,xls_data.manager_working_type)
    set_manager_type(ws2,line,1,resmap)
    ###################################################################
    wb.active = ws1
    wb.save(xls_data.get_filepath())
    savelog(xls_data.get_filepath() + "文件已生成！")

if __name__ == '__main__':
    print("start")
    create_word_table()
