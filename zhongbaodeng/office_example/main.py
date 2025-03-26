import sys
from docx import Document
import docx
from docx.shared import RGBColor
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.oxml import OxmlElement
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from openpyxl.styles import Alignment
from openpyxl.styles import PatternFill
from openpyxl.styles import Alignment
from openpyxl.styles import Border, Side, Font

# 填充表格内容并设置背景颜色
# 第一行设置背景颜色
def set_cell_background(cell, color):
    cell._element.get_or_add_tcPr().append(
        docx.oxml.parse_xml(r'<w:shd {} w:fill="{}"/>'.format(
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"',
            color)
        )
    )

def set_cell_font_color(cell, color):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = color

def set_font(paragraph, font_name, font_size):
    run2 = paragraph.runs[0]
    run2.font.name = font_name
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run2.font.size = Pt(font_size)

def set_text_align(cell,horizontal=True,vertical=True):
    if horizontal:
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if vertical:
        cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

def set_row_height(table, row_idx, height):
    tr = table.rows[row_idx]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

def create_word_table():
    #创建word文件，在其中创建表格
    # 创建一个新的Document对象
    doc = Document()
    # 添加标题
    doc.add_heading('Python创建表格示例', 0)
    doc.add_paragraph('这是一个简单的示例，展示如何使用 python 创建并编辑 Word 文件，以及在word中创建表格。')
    para = doc.add_paragraph()
    runt1 = para.add_run('前面的文字使用默认颜色')
    runt2 = para.add_run('后面的文字使用红色字体')
    runt2.font.color.rgb = RGBColor(255,0,0)
    # 添加一个表格，表格大小为3行4列
    table = doc.add_table(rows=3, cols=5)
    
    set_cell_background(table.cell(0, 0), 'D9D9D9') 
    set_cell_background(table.cell(0, 1), 'D9D9D9')
    set_cell_background(table.cell(0, 2), 'D9D9D9')
    set_cell_background(table.cell(0, 3), 'D9D9D9')
    set_cell_background(table.cell(0, 4), 'D9D9D9')
    table.style = 'Table Grid'
    # 获取表格的单元格并填写内容
    # 填写标题
    cell_00 = table.cell(0, 0)
    cell_00.text = '序\n号'
    set_text_align(cell_00)
    set_font(cell_00.paragraphs[0],'仿宋_GB2312',11)
    
    cell_01 = table.cell(0, 1)
    cell_01.text = '产品名称'
    set_text_align(cell_01)
    set_font(cell_01.paragraphs[0],'仿宋_GB2312',11)

    cell_02 = table.cell(0, 2)
    cell_02.text = '工作进度'
    set_text_align(cell_02)
    set_font(cell_02.paragraphs[0],'仿宋_GB2312',11)
    
    cell_03 = table.cell(0, 3)
    cell_03.text = '产品规模(亿元)'
    set_text_align(cell_03)
    set_font(cell_03.paragraphs[0],'仿宋_GB2312',11)

    cell_04 = table.cell(0, 4)
    cell_04.text = '备注'
    set_text_align(cell_04)
    set_font(cell_04.paragraphs[0],'仿宋_GB2312',11)
    ################填写内容#####################################################
    cell_10 = table.cell(1, 0)
    cell_10.text = 'x'
    set_text_align(cell_10)
    
    cell_11 = table.cell(1, 1)
    paragraph = cell_11.paragraphs[0]
    run1 = paragraph.add_run('xx')
    run1.font.color.rgb = RGBColor(0,0,0)
    run2 = paragraph.add_run('（资产支持计划名称）')
    run2.font.color.rgb = RGBColor(255,0,0)
    set_text_align(cell_11,False,True)

    cell_12 = table.cell(1, 2)
    #cell_12.text = '1.xx年x月x日报送系统，当日分配人员进行查验。\n2.xx年x月x日xx（予以登记/不予登记)。'
    paragraph = cell_12.paragraphs[0]
    run121 = paragraph.add_run('1.xx年x月x日报送系统，当日分配人员进行查验。\n2.xx年x月x日xx')
    run121.font.color.rgb = RGBColor(0,0,0)
    run122 = paragraph.add_run('（予以登记/不予登记）')
    run122.font.color.rgb = RGBColor(255,0,0)
    run123 = paragraph.add_run('。')
    run123.font.color.rgb = RGBColor(0,0,0)
    
    cell_13 = table.cell(1, 3)
    cell_13.text = 'x'
    set_text_align(cell_13)
    
    cell_14 = table.cell(1, 4)
    cell_14.text = '予以登记/不予登记/查验中。'
    set_cell_font_color(cell_14, RGBColor(255, 0, 0)) 
    set_text_align(cell_14,False,True)

    #设置单元格宽度
    for row in table.rows:
        cell0 = row.cells[0]
        cell0.width = Inches(0.5)
        cell2 = row.cells[2]
        cell2.width = Inches(2.5)

    #设置标题栏高度
    set_row_height(table,0,1000)

    # 保存文档
    doc.save('word_table.docx')

    print("Word文件已生成！")

def set_excel_sheet_row_text(ws,line,text_list):
    #设置excel文件某一行的文字
    border_style = Border(
        top=Side(border_style='thin', color='000000'),
        bottom=Side(border_style='thin', color='000000'),
        left=Side(border_style='thin', color='000000'),
        right=Side(border_style='thin', color='000000')
    )
    id = 1
    alphabet=['A','B','C','D','E','F','G','H','I','J','K','L','M','N','O','P','Q','R','S','T','U','V','W','X','Y','Z']
    if len(text_list) > len(alphabet):
        print('set_excel_sheet_row_text text_list and alphabet are not match!')
        sys.exit(1)

    for item in text_list:
        ws.cell(row=line,column=id).value = item
        cell_id = alphabet[id - 1]
        ws.column_dimensions[cell_id].width = 15
        ws[cell_id + str(line)].border = border_style
        if len(item) <= 4:
            ws[cell_id + str(line)].alignment = Alignment(vertical='center')
        else:
            ws[cell_id + str(line)].alignment = Alignment(wrap_text=True)

        id = id + 1

def create_excel_table():
    #创建excel文件内容
    wb = Workbook()
    fill_green = PatternFill(start_color="92D050", end_color="92D050", fill_type="solid") 
    fill_yellow = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
    # 设置颜色
    font_style = Font(color='FF0000')
    
    ws1 = wb.active
    ########################################################################################################################################
    ws1.title = '月度'
    ws1.merge_cells('A1:M1')
    ws1['A1'] = '中保登资产支持计划登记情况表（月度）'
    ws1['A1'].alignment = Alignment(horizontal='center', vertical='center')

    line = 2
    t12 = ["总体情况","xx年x月登记规模","去年x月登记规模","同比增长","xx年x月登记只数","去年x月登记只数","同比增长","xx年x月发行规模","去年x月发行规模","同比增长","xx年x月发行只数","去年x月发行只数","同比增长"]
    set_excel_sheet_row_text(ws1,line,t12)
    ws1['A' + str(line)].fill = fill_green
    ws1['H' + str(line)].fill = fill_yellow
    
    line = 6
    t16 = ["受托人","xx年x月登记规模","去年x月登记规模","同比增长","xx年x月登记只数","去年x月登记只数","同比增长","xx年x月发行规模","去年x月发行规模","同比增长","xx年x月发行只数","去年x月发行只数","同比增长"]
    set_excel_sheet_row_text(ws1,line,t16)
    ws1['A' + str(line)].fill = fill_green
    
    line = 10
    t110 = ["托管银行","xx年x月登记规模","去年x月登记规模","同比增长","xx年x月登记只数","去年x月登记只数","同比增长","xx年x月发行规模","去年x月发行规模","同比增长","xx年x月发行只数","去年x月发行只数","同比增长"]
    set_excel_sheet_row_text(ws1,line,t110)
    ws1['A' + str(line)].fill = fill_green

    #ws1['M2'].font = font_style
    ########################################################################################################################################
    ws2 = wb.create_sheet('季度')
    ws2.merge_cells('A1:M1')
    ws2['A1'] = '中保登资产支持计划登记情况表（季度）'
    ws2['A1'].alignment = Alignment(horizontal='center', vertical='center')

    line = 2
    t22 = ["总体情况","xx年x季度登记规模","去年x季度登记规模","同比增长","xx年x季度登记只数","去年x季度登记只数","同比增长","xx年x季度发行规模","去年x季度发行规模","同比增长","xx年x季度发行只数","去年x季度发行只数","同比增长"]
    set_excel_sheet_row_text(ws2,line,t22)
    ws2['A' + str(line)].fill = fill_green
    ws2['H' + str(line)].fill = fill_yellow

    line = 6
    t26 = ["受托人","xx年x季度登记规模","去年x季度登记规模","同比增长","xx年x季度登记只数","去年x季度登记只数","同比增长","xx年x季度发行规模","去年x季度发行规模","同比增长","xx年x季度发行只数","去年x季度发行只数","同比增长"]
    set_excel_sheet_row_text(ws2,line,t26)
    ws2['A' + str(line)].fill = fill_green

    line = 10
    t210 = ["托管银行","xx年x季度登记规模","去年x季度登记规模","同比增长","xx年x季度登记只数","去年x季度登记只数","同比增长","xx年x季度发行规模","去年x季度发行规模","同比增长","xx年x季度发行只数","去年x季度发行只数","同比增长"]
    set_excel_sheet_row_text(ws2,line,t210)
    ws2['A' + str(line)].fill = fill_green
    ########################################################################################################################################
    ws3 = wb.create_sheet('年度')
    ws3.merge_cells('A1:M1')
    ws3['A1'] = '中保登资产支持计划登记情况表（年度）'
    ws3['A1'].alignment = Alignment(horizontal='center', vertical='center')

    line = 2
    t32 = ["总体情况","xx年累计登记规模","去年累计登记规模","同比增长","xx年累计登记只数","去年累计登记只数","同比增长","xx年累计发行规模","去年累计发行规模","同比增长","xx年累计发行只数","去年累计发行只数","同比增长"]
    set_excel_sheet_row_text(ws3,line,t32)
    ws3['A' + str(line)].fill = fill_green
    ws3['H' + str(line)].fill = fill_yellow

    line = 6
    t36 = ["受托人","xx年累计登记规模","去年累计登记规模","同比增长","xx年累计登记只数","去年累计登记只数","同比增长","xx年累计发行规模","去年累计发行规模","同比增长","xx年累计发行只数","去年累计发行只数","同比增长"]
    set_excel_sheet_row_text(ws3,line,t36)
    ws3['A' + str(line)].fill = fill_green

    line = 10
    t310 = ["托管银行","xx年累计登记规模","去年累计登记规模","同比增长","xx年累计登记只数","去年累计登记只数","同比增长","xx年累计发行规模","去年累计发行规模","同比增长","xx年累计发行只数","去年累计发行只数","同比增长"]
    set_excel_sheet_row_text(ws3,line,t310)
    ws3['A' + str(line)].fill = fill_green
    ########################################################################################################################################
    ws4 = wb.create_sheet('历史累计')
    ws4.merge_cells('A1:M1')
    ws4['A1'] = '中保登资产支持计划登记情况表（历史累计）'
    ws4['A1'].alignment = Alignment(horizontal='center', vertical='center')

    line = 2
    t42 = ["总体情况","历史累计登记规模（截至xx年x月x日）","历史累计登记规模（截至去年x月x日）","同比增长","历史累计登记只数（截至xx年x月x日）","历史累计登记只数（截至去年x月x日）","同比增长","存量规模（截至xx年x月x日）","去年存量规模（截至xx年x月x日）","同比增长"]
    set_excel_sheet_row_text(ws4,line,t42)
    ws4['A' + str(line)].fill = fill_green
    ws4['H' + str(line)].fill = fill_yellow

    line = 6
    t46 = ["受托人","历史累计登记规模（截至xx年x月x日）","历史累计登记规模（截至去年x月x日）","同比增长","历史累计登记只数（截至xx年x月x日）","历史累计登记只数（截至去年x月x日）","同比增长","存量规模（截至xx年x月x日）","去年存量规模","同比增长"]
    set_excel_sheet_row_text(ws4,line,t46)
    ws4['A' + str(line)].fill = fill_green

    line = 10
    t410 = ["托管银行","历史累计登记规模（截至xx年x月x日）","历史累计登记规模（截至去年x月x日）","同比增长","历史累计登记只数（截至xx年x月x日）","历史累计登记只数（截至去年x月x日）","同比增长","存量规模（截至xx年x月x日）","去年存量规模","同比增长"]
    set_excel_sheet_row_text(ws4,line,t410)
    ws4['A' + str(line)].fill = fill_green
    ########################################################################################################################################
    wb.active = ws1
    wb.save('excel.xlsx')
    print("Excel文件已生成！")


if __name__ == '__main__':
    print("start")
    create_word_table()
    create_excel_table()
    
